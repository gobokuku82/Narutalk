"""
Word 문서 읽기 테스트 - 인코딩 문제 해결
"""

import docx
from pathlib import Path
import sys

def test_read_docx():
    """DOCX 파일 읽기 테스트"""

    doc_path = Path(r"C:\kdy\Projects\Narutalk_V003\beta_v0031\database\Rule_DB\의약품 리베이트 및 광고, 지출보고서 관련 법령 및 규약 통합본.docx")

    print(f"문서 경로: {doc_path}")
    print(f"파일 존재: {doc_path.exists()}")

    if not doc_path.exists():
        print("ERROR: 파일이 존재하지 않습니다.")
        return

    try:
        # 문서 열기
        doc = docx.Document(str(doc_path))

        print(f"\n총 단락 수: {len(doc.paragraphs)}")
        print(f"총 표 수: {len(doc.tables)}")

        # 처음 10개 단락 출력
        print("\n=== 처음 10개 단락 ===")
        for i, para in enumerate(doc.paragraphs[:10]):
            text = para.text.strip()
            if text:
                # Windows 콘솔 인코딩 처리
                try:
                    print(f"단락 {i+1}: {text[:100]}...")
                except UnicodeEncodeError:
                    # cp949로 인코딩할 수 없는 문자 처리
                    safe_text = text.encode('cp949', errors='replace').decode('cp949')
                    print(f"단락 {i+1}: {safe_text[:100]}...")

        # 전체 텍스트 추출
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 주요 키워드 검색
        print("\n=== 키워드 검색 ===")
        keywords = ['청탁금지법', '공정경쟁규약', '약사법', '제품설명회', '학술대회', '견본품', '강연료']

        combined_text = '\n'.join(full_text)
        for keyword in keywords:
            if keyword in combined_text:
                print(f"[Found] '{keyword}'")
                # 해당 키워드가 포함된 첫 문장 찾기
                for text in full_text:
                    if keyword in text:
                        try:
                            print(f"  Example: {text[:80]}...")
                        except UnicodeEncodeError:
                            safe_text = text.encode('cp949', errors='replace').decode('cp949')
                            print(f"  Example: {safe_text[:80]}...")
                        break

        # 파일에 저장 (UTF-8)
        output_path = Path("test_document_content.txt")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        print(f"\n전체 내용을 {output_path}에 저장했습니다.")

        return combined_text

    except Exception as e:
        print(f"ERROR: 문서 읽기 실패 - {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # Windows 콘솔 UTF-8 설정
    if sys.platform == 'win32':
        import os
        os.system('chcp 65001 > nul')

    test_read_docx()