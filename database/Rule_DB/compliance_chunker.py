"""
제약 영업 규제 준수 시스템 - 문서 청킹 모듈
"""

import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
import docx
from enum import Enum


class ProhibitionType(Enum):
    ABSOLUTE = "절대금지"
    CONDITIONAL = "조건부허용"
    ALLOWED = "허용"


class LawType(Enum):
    ANTI_GRAFT = "청탁금지법"
    PHARMACEUTICAL = "약사법"
    FAIR_COMPETITION = "공정경쟁규약"
    FAIR_TRADE = "공정거래법"
    GUIDELINE = "가이드라인"


@dataclass
class ChunkMetadata:
    law_name: str
    article: Optional[str] = None
    paragraph: Optional[str] = None
    subparagraph: Optional[str] = None
    prohibition_type: Optional[str] = None
    limit_value: Optional[int] = None
    limit_unit: Optional[str] = None
    frequency_count: Optional[int] = None
    frequency_period: Optional[str] = None
    target: Optional[str] = None
    activity: Optional[str] = None
    item_type: Optional[str] = None
    purpose: Optional[str] = None
    conditions: Optional[List[str]] = None
    exceptions: Optional[List[str]] = None
    keywords: Optional[List[str]] = None
    effective_date: Optional[str] = None
    priority: Optional[int] = None


@dataclass
class Chunk:
    chunk_id: str
    text: str
    metadata: ChunkMetadata
    source_location: Optional[Dict[str, Any]] = None
    created_at: str = None

    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now().isoformat()


class ComplianceChunker:
    def __init__(self, doc_path: str):
        """
        문서 로드 및 초기화

        Args:
            doc_path: DOCX 문서 경로
        """
        self.doc_path = Path(doc_path)
        self.chunks: List[Chunk] = []
        self.chunk_counter = 0

        # 청킹 설정
        self.min_chunk_size = 50
        self.max_chunk_size = 300
        self.optimal_chunk_size = 150
        self.overlap_size = 30

        # 법령 우선순위
        self.law_priority = {
            LawType.ANTI_GRAFT.value: 1,
            LawType.PHARMACEUTICAL.value: 2,
            LawType.FAIR_TRADE.value: 3,
            LawType.FAIR_COMPETITION.value: 4,
            LawType.GUIDELINE.value: 5
        }

        # 패턴 정의
        self.patterns = {
            'amount': [
                (r'(\d+)만\s*원', 10000),
                (r'(\d+)천\s*원', 1000),
                (r'(\d+)백만\s*원', 1000000),
                (r'(\d+)원', 1)
            ],
            'frequency': [
                (r'월\s*(\d+)회', 'month'),
                (r'연\s*(\d+)회', 'year'),
                (r'일\s*(\d+)회', 'day'),
                (r'주\s*(\d+)회', 'week')
            ],
            'article': r'제(\d+)조',
            'paragraph': r'제(\d+)항',
            'subparagraph': r'제(\d+)호',
            'exception': r'(다만|단,|단서|예외)',
            'target': r'(의료인|의료기관|요양기관|공직자|보건의료전문가)',
            'activity': r'(제품설명회|학술대회|강연|자문|임상시험|시판후조사)'
        }

    def load_document(self) -> str:
        """DOCX 문서 로드 및 텍스트 추출"""
        doc = docx.Document(str(self.doc_path))
        full_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())

        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        return '\n'.join(full_text)

    def chunk_by_law(self, text: str, law_type: str) -> List[Chunk]:
        """
        법령별 청킹 전략 적용

        Args:
            text: 청킹할 텍스트
            law_type: 법령 유형

        Returns:
            청킹된 결과 리스트
        """
        chunks = []

        if law_type == LawType.ANTI_GRAFT.value:
            chunks.extend(self._chunk_anti_graft_law(text))
        elif law_type == LawType.FAIR_COMPETITION.value:
            chunks.extend(self._chunk_fair_competition(text))
        elif law_type == LawType.PHARMACEUTICAL.value:
            chunks.extend(self._chunk_pharmaceutical_law(text))
        else:
            chunks.extend(self._chunk_generic(text, law_type))

        return chunks

    def _chunk_anti_graft_law(self, text: str) -> List[Chunk]:
        """청탁금지법 특화 청킹"""
        chunks = []

        # 조항별로 분리
        article_pattern = re.compile(r'제(\d+)조[^제]*?(?=제\d+조|$)', re.DOTALL)
        articles = article_pattern.findall(text)

        for article_match in article_pattern.finditer(text):
            article_num = article_match.group(1)
            article_text = article_match.group(0)

            # 금액 제한 추출
            for pattern, multiplier in self.patterns['amount']:
                for match in re.finditer(pattern, article_text):
                    amount = int(match.group(1)) * multiplier
                    context = self._extract_context(article_text, match.span())

                    # 금지/허용 유형 판단
                    prohibition_type = self._determine_prohibition_type(context)

                    # 메타데이터 생성
                    metadata = ChunkMetadata(
                        law_name=LawType.ANTI_GRAFT.value,
                        article=f"제{article_num}조",
                        prohibition_type=prohibition_type,
                        limit_value=amount,
                        priority=self.law_priority[LawType.ANTI_GRAFT.value]
                    )

                    # 추가 메타데이터 추출
                    metadata = self._extract_additional_metadata(context, metadata)

                    chunk = Chunk(
                        chunk_id=f"anti_graft_{article_num}_{self.chunk_counter}",
                        text=context,
                        metadata=metadata
                    )
                    chunks.append(chunk)
                    self.chunk_counter += 1

        return chunks

    def _chunk_fair_competition(self, text: str) -> List[Chunk]:
        """공정경쟁규약 특화 청킹"""
        chunks = []

        # 활동별 패턴 정의
        activity_patterns = {
            '제품설명회': r'제품설명회[^.]*?[.]',
            '학술대회': r'학술대회[^.]*?[.]',
            '견본품': r'견본품[^.]*?[.]',
            '강연료': r'강연료[^.]*?[.]'
        }

        for activity, pattern in activity_patterns.items():
            for match in re.finditer(pattern, text, re.DOTALL):
                context = match.group(0)

                # 메타데이터 생성
                metadata = ChunkMetadata(
                    law_name=LawType.FAIR_COMPETITION.value,
                    activity=activity,
                    priority=self.law_priority[LawType.FAIR_COMPETITION.value]
                )

                # 금액 추출
                for amount_pattern, multiplier in self.patterns['amount']:
                    amount_match = re.search(amount_pattern, context)
                    if amount_match:
                        metadata.limit_value = int(amount_match.group(1)) * multiplier
                        break

                # 빈도 추출
                for freq_pattern, period in self.patterns['frequency']:
                    freq_match = re.search(freq_pattern, context)
                    if freq_match:
                        metadata.frequency_count = int(freq_match.group(1))
                        metadata.frequency_period = period
                        break

                # 추가 메타데이터
                metadata = self._extract_additional_metadata(context, metadata)

                chunk = Chunk(
                    chunk_id=f"fair_comp_{activity}_{self.chunk_counter}",
                    text=context,
                    metadata=metadata
                )
                chunks.append(chunk)
                self.chunk_counter += 1

        return chunks

    def _chunk_pharmaceutical_law(self, text: str) -> List[Chunk]:
        """약사법 특화 청킹"""
        chunks = []

        # 리베이트 관련 조항 중심 청킹
        rebate_patterns = [
            r'리베이트[^.]*?[.]',
            r'경제적 이익[^.]*?[.]',
            r'금품[^.]*?[.]'
        ]

        for pattern in rebate_patterns:
            for match in re.finditer(pattern, text, re.DOTALL):
                context = match.group(0)

                metadata = ChunkMetadata(
                    law_name=LawType.PHARMACEUTICAL.value,
                    priority=self.law_priority[LawType.PHARMACEUTICAL.value]
                )

                # 메타데이터 추출
                metadata = self._extract_additional_metadata(context, metadata)

                chunk = Chunk(
                    chunk_id=f"pharma_{self.chunk_counter}",
                    text=context,
                    metadata=metadata
                )
                chunks.append(chunk)
                self.chunk_counter += 1

        return chunks

    def _chunk_generic(self, text: str, law_type: str) -> List[Chunk]:
        """일반적인 청킹 전략"""
        chunks = []
        sentences = re.split(r'[.]\s+', text)

        current_chunk = []
        current_size = 0

        for sentence in sentences:
            sentence_size = len(sentence)

            if current_size + sentence_size > self.max_chunk_size and current_chunk:
                # 현재 청크 저장
                chunk_text = '. '.join(current_chunk) + '.'
                metadata = ChunkMetadata(
                    law_name=law_type,
                    priority=self.law_priority.get(law_type, 99)
                )
                metadata = self._extract_additional_metadata(chunk_text, metadata)

                chunk = Chunk(
                    chunk_id=f"generic_{self.chunk_counter}",
                    text=chunk_text,
                    metadata=metadata
                )
                chunks.append(chunk)
                self.chunk_counter += 1

                # 오버랩 처리
                if self.overlap_size > 0 and len(current_chunk) > 1:
                    current_chunk = current_chunk[-2:]
                    current_size = sum(len(s) for s in current_chunk)
                else:
                    current_chunk = []
                    current_size = 0

            current_chunk.append(sentence)
            current_size += sentence_size

        # 마지막 청크 처리
        if current_chunk:
            chunk_text = '. '.join(current_chunk) + '.'
            metadata = ChunkMetadata(
                law_name=law_type,
                priority=self.law_priority.get(law_type, 99)
            )
            metadata = self._extract_additional_metadata(chunk_text, metadata)

            chunk = Chunk(
                chunk_id=f"generic_{self.chunk_counter}",
                text=chunk_text,
                metadata=metadata
            )
            chunks.append(chunk)
            self.chunk_counter += 1

        return chunks

    def _extract_context(self, text: str, span: Tuple[int, int], window: int = 100) -> str:
        """매치된 부분 주변 컨텍스트 추출"""
        start = max(0, span[0] - window)
        end = min(len(text), span[1] + window)

        context = text[start:end]

        # 문장 경계 찾기
        first_period = context.find('.')
        last_period = context.rfind('.')

        if first_period != -1 and last_period != -1 and first_period < last_period:
            context = context[first_period+1:last_period+1]

        return context.strip()

    def _determine_prohibition_type(self, text: str) -> str:
        """텍스트에서 금지 유형 판단"""
        if any(word in text for word in ['금지', '불가', '안 된다', '할 수 없']):
            return ProhibitionType.ABSOLUTE.value
        elif any(word in text for word in ['범위 내', '이내', '까지', '허용']):
            return ProhibitionType.CONDITIONAL.value
        else:
            return ProhibitionType.ALLOWED.value

    def _extract_additional_metadata(self, text: str, metadata: ChunkMetadata) -> ChunkMetadata:
        """텍스트에서 추가 메타데이터 추출"""

        # 대상 추출
        target_match = re.search(self.patterns['target'], text)
        if target_match:
            metadata.target = target_match.group(1)

        # 활동 추출
        if not metadata.activity:
            activity_match = re.search(self.patterns['activity'], text)
            if activity_match:
                metadata.activity = activity_match.group(1)

        # 예외 조항 확인
        if re.search(self.patterns['exception'], text):
            metadata.exceptions = [text[match.start():match.end()+50]
                                 for match in re.finditer(self.patterns['exception'], text)]

        # 키워드 추출
        keywords = []
        if '식사' in text or '식음료' in text:
            keywords.append('식음료')
        if '숙박' in text:
            keywords.append('숙박')
        if '교통' in text:
            keywords.append('교통비')
        if '등록비' in text or '참가비' in text:
            keywords.append('참가비')

        if keywords:
            metadata.keywords = keywords

        return metadata

    def extract_metadata(self, text: str) -> Dict[str, Any]:
        """텍스트에서 메타데이터 자동 추출"""
        metadata = {}

        # 조항 정보
        article_match = re.search(self.patterns['article'], text)
        if article_match:
            metadata['article'] = f"제{article_match.group(1)}조"

        paragraph_match = re.search(self.patterns['paragraph'], text)
        if paragraph_match:
            metadata['paragraph'] = f"제{paragraph_match.group(1)}항"

        # 금액 정보
        for pattern, multiplier in self.patterns['amount']:
            match = re.search(pattern, text)
            if match:
                metadata['limit_value'] = int(match.group(1)) * multiplier
                break

        # 빈도 정보
        for pattern, period in self.patterns['frequency']:
            match = re.search(pattern, text)
            if match:
                metadata['frequency_count'] = int(match.group(1))
                metadata['frequency_period'] = period
                break

        return metadata

    def process_document(self) -> List[Chunk]:
        """전체 문서 처리 및 청킹"""
        print(f"문서 로드 중: {self.doc_path}")
        full_text = self.load_document()

        # 법령별로 섹션 분리 (간단한 휴리스틱)
        sections = self._split_by_law_sections(full_text)

        all_chunks = []
        for law_type, section_text in sections.items():
            print(f"{law_type} 청킹 중...")
            chunks = self.chunk_by_law(section_text, law_type)
            all_chunks.extend(chunks)

        self.chunks = all_chunks
        return all_chunks

    def _split_by_law_sections(self, text: str) -> Dict[str, str]:
        """텍스트를 법령별 섹션으로 분리"""
        sections = {}

        # 법령 구분 패턴
        law_markers = {
            LawType.ANTI_GRAFT.value: ['청탁금지법', '부정청탁', '금품등'],
            LawType.PHARMACEUTICAL.value: ['약사법', '의약품', '리베이트'],
            LawType.FAIR_COMPETITION.value: ['공정경쟁규약', '공정거래', '의약품 거래'],
        }

        # 간단한 분류 (실제로는 더 정교한 로직 필요)
        current_section = []
        current_law = None

        lines = text.split('\n')
        for line in lines:
            # 법령 마커 확인
            detected_law = None
            for law, markers in law_markers.items():
                if any(marker in line for marker in markers):
                    detected_law = law
                    break

            if detected_law and detected_law != current_law:
                # 이전 섹션 저장
                if current_law and current_section:
                    sections[current_law] = '\n'.join(current_section)
                # 새 섹션 시작
                current_law = detected_law
                current_section = [line]
            elif current_law:
                current_section.append(line)

        # 마지막 섹션 저장
        if current_law and current_section:
            sections[current_law] = '\n'.join(current_section)

        # 전체 텍스트를 일반 섹션으로 추가 (분류되지 않은 부분)
        if not sections:
            sections['일반'] = text

        return sections

    def save_chunks(self, output_path: str = None):
        """청킹 결과를 JSON 파일로 저장"""
        if output_path is None:
            output_path = self.doc_path.parent / 'chunked_data.json'

        output_data = {
            'metadata': {
                'source_document': str(self.doc_path),
                'total_chunks': len(self.chunks),
                'created_at': datetime.now().isoformat(),
                'chunking_params': {
                    'min_size': self.min_chunk_size,
                    'max_size': self.max_chunk_size,
                    'overlap': self.overlap_size
                }
            },
            'chunks': [
                {
                    'chunk_id': chunk.chunk_id,
                    'text': chunk.text,
                    'metadata': asdict(chunk.metadata),
                    'created_at': chunk.created_at
                }
                for chunk in self.chunks
            ]
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)

        print(f"청킹 결과 저장 완료: {output_path}")
        print(f"총 {len(self.chunks)}개 청크 생성")

        return output_path

    def validate_chunks(self) -> Dict[str, Any]:
        """청킹 품질 검증"""
        validation_results = {
            'total_chunks': len(self.chunks),
            'size_distribution': {},
            'metadata_completeness': {},
            'issues': []
        }

        # 크기 분포 분석
        sizes = [len(chunk.text) for chunk in self.chunks]
        validation_results['size_distribution'] = {
            'min': min(sizes) if sizes else 0,
            'max': max(sizes) if sizes else 0,
            'avg': sum(sizes) / len(sizes) if sizes else 0,
            'within_range': sum(1 for s in sizes if self.min_chunk_size <= s <= self.max_chunk_size)
        }

        # 메타데이터 완성도
        metadata_fields = ['law_name', 'prohibition_type', 'limit_value', 'target', 'activity']
        field_counts = {field: 0 for field in metadata_fields}

        for chunk in self.chunks:
            for field in metadata_fields:
                if getattr(chunk.metadata, field) is not None:
                    field_counts[field] += 1

        validation_results['metadata_completeness'] = {
            field: f"{(count/len(self.chunks)*100):.1f}%" if self.chunks else "0%"
            for field, count in field_counts.items()
        }

        # 문제점 확인
        for chunk in self.chunks:
            if len(chunk.text) < self.min_chunk_size:
                validation_results['issues'].append(f"청크 {chunk.chunk_id}: 크기가 너무 작음 ({len(chunk.text)}자)")
            elif len(chunk.text) > self.max_chunk_size:
                validation_results['issues'].append(f"청크 {chunk.chunk_id}: 크기가 너무 큼 ({len(chunk.text)}자)")

            if not chunk.metadata.law_name:
                validation_results['issues'].append(f"청크 {chunk.chunk_id}: 법령명 누락")

        return validation_results


def main():
    """테스트 실행"""
    doc_path = r"C:\kdy\Projects\Narutalk_V003\beta_v0031\database\Rule_DB\의약품 리베이트 및 광고, 지출보고서 관련 법령 및 규약 통합본.docx"

    chunker = ComplianceChunker(doc_path)
    chunks = chunker.process_document()

    # 결과 저장
    output_path = chunker.save_chunks()

    # 검증
    validation = chunker.validate_chunks()
    print("\n=== 청킹 검증 결과 ===")
    print(json.dumps(validation, ensure_ascii=False, indent=2))

    # 샘플 출력
    print("\n=== 샘플 청크 (처음 3개) ===")
    for chunk in chunks[:3]:
        print(f"\nID: {chunk.chunk_id}")
        print(f"Text: {chunk.text[:100]}...")
        print(f"Metadata: {json.dumps(asdict(chunk.metadata), ensure_ascii=False, indent=2)}")


if __name__ == "__main__":
    main()