"""
DOCX 파일 로더 - 좋은제약 내부 규정 문서 처리
한글 인코딩 문제 해결 및 구조화된 텍스트 추출
"""

import docx
from pathlib import Path
import sys
import os
from typing import List, Dict, Optional
import re

class DocxLoader:
    """DOCX 문서 로더 및 파서"""

    def __init__(self, file_path: str):
        """
        Args:
            file_path: DOCX 파일 경로
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        self.doc = None
        self.paragraphs = []
        self.tables = []

        # Windows 콘솔 UTF-8 설정
        if sys.platform == 'win32':
            os.system('chcp 65001 > nul 2>&1')
            # 표준 출력 인코딩 설정
            sys.stdout.reconfigure(encoding='utf-8')

    def load_document(self) -> bool:
        """문서 로드"""
        try:
            self.doc = docx.Document(str(self.file_path))
            print(f"[OK] 문서 로드 성공: {self.file_path.name}")
            print(f"  - 단락 수: {len(self.doc.paragraphs)}")
            print(f"  - 표 수: {len(self.doc.tables)}")
            return True
        except Exception as e:
            print(f"[ERROR] 문서 로드 실패: {e}")
            return False

    def extract_paragraphs(self) -> List[Dict[str, any]]:
        """단락 추출 및 구조화"""
        if not self.doc:
            self.load_document()

        self.paragraphs = []
        current_part = None
        current_chapter = None

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            para_info = {
                'index': i,
                'text': text,
                'style': para.style.name if para.style else None,
                'part': current_part,
                'chapter': current_chapter,
                'article': None,
                'type': 'text'
            }

            # 부(Part) 인식
            if '제1부' in text and '취업규칙' in text:
                current_part = '제1부_취업규칙'
                para_info['type'] = 'part_header'
            elif '제2부' in text and '인사관리' in text:
                current_part = '제2부_인사관리규정'
                para_info['type'] = 'part_header'
            elif '제3부' in text and '복무규정' in text:
                current_part = '제3부_복무규정'
                para_info['type'] = 'part_header'
            elif '제4부' in text and ('윤리' in text or 'CP' in text):
                current_part = '제4부_윤리강령_CP'
                para_info['type'] = 'part_header'

            # 장(Chapter) 인식
            chapter_match = re.search(r'제(\d+)장\s*(.+)', text)
            if chapter_match:
                current_chapter = f"제{chapter_match.group(1)}장"
                para_info['chapter_title'] = chapter_match.group(2).strip()
                para_info['type'] = 'chapter_header'

            # 조(Article) 인식
            article_match = re.search(r'제(\d+)조\s*[\(（]([^)）]+)[\)）]', text)
            if article_match:
                para_info['article'] = f"제{article_match.group(1)}조"
                para_info['article_title'] = article_match.group(2)
                para_info['type'] = 'article'

            # 항(Clause) 인식
            clause_match = re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text)
            if clause_match:
                para_info['type'] = 'clause'

            # 번호 리스트 인식
            number_match = re.match(r'^(\d+)\.\s+', text)
            if number_match:
                para_info['type'] = 'numbered_list'

            self.paragraphs.append(para_info)

        return self.paragraphs

    def extract_tables(self) -> List[Dict[str, any]]:
        """표 데이터 추출"""
        if not self.doc:
            self.load_document()

        self.tables = []

        for i, table in enumerate(self.doc.tables):
            table_data = {
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': [],
                'text_representation': ''
            }

            # 표 데이터를 2차원 리스트로 추출
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data['data'].append(row_data)

            # 텍스트 표현 생성 (구조화된 텍스트)
            text_lines = []
            for row_idx, row in enumerate(table_data['data']):
                if row_idx == 0:  # 헤더로 가정
                    text_lines.append(' | '.join(row))
                    text_lines.append('-' * 50)
                else:
                    text_lines.append(' | '.join(row))

            table_data['text_representation'] = '\n'.join(text_lines)
            self.tables.append(table_data)

        return self.tables

    def get_full_text(self) -> str:
        """전체 텍스트 추출"""
        if not self.paragraphs:
            self.extract_paragraphs()

        full_text = []
        for para in self.paragraphs:
            full_text.append(para['text'])

        # 표 텍스트 추가
        if not self.tables:
            self.extract_tables()

        for table in self.tables:
            full_text.append('\n[표]\n')
            full_text.append(table['text_representation'])
            full_text.append('\n')

        return '\n'.join(full_text)

    def get_structured_content(self) -> Dict[str, any]:
        """구조화된 콘텐츠 반환"""
        if not self.paragraphs:
            self.extract_paragraphs()
        if not self.tables:
            self.extract_tables()

        # 부별로 그룹화
        structured = {
            '제1부_취업규칙': [],
            '제2부_인사관리규정': [],
            '제3부_복무규정': [],
            '제4부_윤리강령_CP': [],
            'tables': self.tables
        }

        for para in self.paragraphs:
            if para['part']:
                structured[para['part']].append(para)

        return structured

    def search_keywords(self, keywords: List[str]) -> Dict[str, List[str]]:
        """키워드 검색"""
        if not self.paragraphs:
            self.extract_paragraphs()

        results = {keyword: [] for keyword in keywords}

        for para in self.paragraphs:
            text = para['text']
            for keyword in keywords:
                if keyword in text:
                    results[keyword].append({
                        'text': text[:200] + '...' if len(text) > 200 else text,
                        'part': para['part'],
                        'article': para['article']
                    })

        return results

    def save_to_file(self, output_path: str = None):
        """추출된 텍스트를 파일로 저장"""
        if output_path is None:
            output_path = self.file_path.stem + '_extracted.txt'

        output_path = Path(output_path)
        full_text = self.get_full_text()

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_text)

        print(f"[OK] 텍스트 저장 완료: {output_path}")
        return output_path


def test_docx_loader():
    """DOCX 로더 테스트"""
    print("="*60)
    print("DOCX 파일 로더 테스트")
    print("="*60)

    # 파일 경로
    doc_path = r"C:\kdy\Projects\Narutalk_V003\beta_v0031\database\raw_data\vector_db\좋은제약 내부 규정.docx"

    try:
        # 로더 초기화
        loader = DocxLoader(doc_path)

        # 문서 로드
        if loader.load_document():
            # 단락 추출
            print("\n1. 단락 추출 중...")
            paragraphs = loader.extract_paragraphs()
            print(f"   추출된 단락: {len(paragraphs)}개")

            # 처음 5개 단락 출력
            print("\n2. 샘플 단락:")
            for para in paragraphs[:5]:
                if para['text']:
                    print(f"   [{para['type']}] {para['text'][:80]}...")

            # 표 추출
            print("\n3. 표 데이터 추출 중...")
            tables = loader.extract_tables()
            print(f"   추출된 표: {len(tables)}개")

            if tables:
                print(f"   첫 번째 표: {tables[0]['rows']}행 x {tables[0]['cols']}열")

            # 구조화된 콘텐츠
            print("\n4. 구조화된 콘텐츠 분석...")
            structured = loader.get_structured_content()
            for part_name, content in structured.items():
                if part_name != 'tables':
                    print(f"   {part_name}: {len(content)}개 단락")

            # 키워드 검색
            print("\n5. 키워드 검색 테스트...")
            keywords = ['윤리', '징계', '휴가', '평가', '보안']
            search_results = loader.search_keywords(keywords)

            for keyword, results in search_results.items():
                if results:
                    print(f"   '{keyword}': {len(results)}개 발견")

            # 파일 저장
            print("\n6. 추출된 텍스트 저장...")
            output_file = loader.save_to_file('internal_regulations_extracted.txt')

            print("\n[OK] DOCX 로더 테스트 완료!")
            return loader

    except Exception as e:
        print(f"\n[ERROR] 오류 발생: {e}")
        import traceback
        traceback.print_exc()
        return None


if __name__ == "__main__":
    test_docx_loader()